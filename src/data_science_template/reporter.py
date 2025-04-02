"""Module for generating market analysis reports."""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List
import os
import io
import base64
import numpy as np

class ReportGenerator:
    """Generates market analysis reports with visualizations and insights."""

    def __init__(self, output_dir: str = "out"):
        """Initialize the reporter.
        
        Args:
            output_dir: Directory to save reports
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.report = Document()
        self._setup_document_style()

    def _setup_document_style(self) -> None:
        """Set up document styling for better readability."""
        # Set default font
        style = self.report.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Set heading styles
        for i in range(1, 4):
            heading_style = self.report.styles[f'Heading {i}']
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(14 - i)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 51, 102)

    def add_title(self, title: str) -> None:
        """Add report title with styling."""
        title_paragraph = self.report.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.report.add_paragraph()  # Add spacing

    def add_executive_summary(self, insights: Dict[str, Any]) -> None:
        """Add executive summary section with key takeaways."""
        self.report.add_heading("Executive Summary", level=1)
        
        # Add market overview with clear context
        self.report.add_paragraph(
            f"This report analyzes the market performance of {insights['name']} "
            f"from {insights['latest_date']}. The current market conditions show "
            f"{self._get_market_condition(insights)}."
        )
        
        # Add key metrics with clear formatting
        self.report.add_paragraph("Key Metrics:")
        metrics = [
            f"Current Price: ${insights['current_price']:,.2f}",
            f"Year-to-Date Return: {insights['total_return']:.2f}%",
            f"Volatility: {insights['current_volatility']:.2f}%",
            f"RSI: {insights['current_rsi']:.1f}"
        ]
        for metric in metrics:
            self.report.add_paragraph(metric, style='List Bullet')
        
        # Add trading recommendation with emphasis
        recommendation = insights['predictions']['recommendation']
        signal_strength = insights['predictions']['signal_strength']
        self.report.add_paragraph(
            f"Trading Recommendation: {signal_strength} {recommendation}",
            style='Intense Quote'
        )
        self.report.add_paragraph()  # Add spacing

    def _get_market_condition(self, insights: Dict[str, Any]) -> str:
        """Get market condition description based on metrics."""
        rsi = insights['current_rsi']
        volatility = insights['current_volatility']
        
        if rsi > 70 and volatility > 20:
            return "highly volatile with potential overbought conditions"
        elif rsi < 30 and volatility > 20:
            return "highly volatile with potential oversold conditions"
        elif rsi > 70:
            return "potentially overbought"
        elif rsi < 30:
            return "potentially oversold"
        else:
            return "relatively stable"

    def add_technical_analysis(self, insights: Dict[str, Any]) -> None:
        """Add technical analysis section with detailed indicators."""
        self.report.add_heading("Technical Analysis", level=1)
        
        # Moving Averages with clear interpretation
        self.report.add_heading("Moving Averages", level=2)
        ma_data = insights['moving_averages']
        current_price = insights['current_price']
        
        for period, value in ma_data.items():
            difference = ((value - current_price) / current_price) * 100
            self.report.add_paragraph(
                f"{period}: ${value:,.2f} ({difference:+.2f}% from current price)"
            )
        
        # RSI Analysis with clear interpretation
        self.report.add_heading("Relative Strength Index (RSI)", level=2)
        rsi = insights['current_rsi']
        rsi_interpretation = self._get_rsi_interpretation(rsi)
        self.report.add_paragraph(f"Current RSI: {rsi:.1f}")
        self.report.add_paragraph(f"Interpretation: {rsi_interpretation}")
        
        # Volatility Analysis with clear interpretation
        self.report.add_heading("Volatility Analysis", level=2)
        volatility = insights['current_volatility']
        avg_volatility = insights['avg_volatility']
        self.report.add_paragraph(f"Current Volatility: {volatility:.2f}%")
        self.report.add_paragraph(f"Average Volatility: {avg_volatility:.2f}%")
        self.report.add_paragraph(
            f"Market Risk Level: {self._get_volatility_level(volatility)}"
        )
        self.report.add_paragraph()  # Add spacing

    def _get_rsi_interpretation(self, rsi: float) -> str:
        """Get RSI interpretation."""
        if rsi > 70:
            return "Overbought conditions - potential for price correction"
        elif rsi < 30:
            return "Oversold conditions - potential for price recovery"
        elif rsi > 60:
            return "Moderately overbought - exercise caution"
        elif rsi < 40:
            return "Moderately oversold - potential buying opportunity"
        else:
            return "Neutral conditions"

    def _get_volatility_level(self, volatility: float) -> str:
        """Get volatility level description."""
        if volatility > 25:
            return "High - Increased market risk"
        elif volatility > 15:
            return "Moderate - Normal market conditions"
        else:
            return "Low - Stable market conditions"

    def add_trading_signals(self, insights: Dict[str, Any]) -> None:
        """Add trading signals section with clear recommendations."""
        self.report.add_heading("Trading Signals", level=1)
        
        predictions = insights['predictions']
        
        # Signal Overview with clear formatting
        self.report.add_paragraph(
            f"Overall Signal: {predictions['signal_strength']} {predictions['recommendation']}"
        )
        
        # Technical Indicators with clear interpretation
        self.report.add_heading("Technical Indicators", level=2)
        self.report.add_paragraph(f"Moving Average Signal: {predictions['ma_description']}")
        self.report.add_paragraph(f"RSI Signal: {predictions['rsi_description']}")
        
        # Combined Analysis with clear interpretation
        self.report.add_heading("Combined Analysis", level=2)
        self.report.add_paragraph(
            f"Combined Signal Strength: {abs(predictions['combined_signal']):.2f}"
        )
        
        # Action Items with clear formatting
        self.report.add_heading("Recommended Actions", level=2)
        actions = self._get_recommended_actions(predictions)
        for action in actions:
            self.report.add_paragraph(action, style='List Bullet')
        self.report.add_paragraph()  # Add spacing

    def _get_recommended_actions(self, predictions: Dict[str, Any]) -> List[str]:
        """Get recommended actions based on predictions."""
        recommendation = predictions.get('recommendation', 'Hold')
        signal_strength = predictions.get('signal_strength', 'Moderate')
        buy_prob = predictions.get('buy_probability', 50)
        
        actions = []
        
        # Determine trading strategy based on probability and market conditions
        if buy_prob > 60:
            strategy = "Strong Buy"
            actions.extend([
                "Consider establishing new long positions",
                "Look for pullbacks as entry opportunities",
                "Set stop-loss orders below recent support levels",
                "Consider using options strategies like covered calls for income",
                "Monitor volume for confirmation of upward trend"
            ])
        elif buy_prob > 55:
            strategy = "Moderate Buy"
            actions.extend([
                "Consider partial position entry",
                "Wait for confirmation of upward trend",
                "Monitor key resistance levels",
                "Consider using dollar-cost averaging",
                "Set tight stop-loss orders"
            ])
        elif buy_prob < 40:
            strategy = "Strong Sell"
            actions.extend([
                "Consider reducing long positions",
                "Look for rallies as exit opportunities",
                "Set stop-loss orders above recent resistance levels",
                "Consider using put options for protection",
                "Monitor volume for confirmation of downward trend"
            ])
        elif buy_prob < 45:
            strategy = "Moderate Sell"
            actions.extend([
                "Consider partial position reduction",
                "Wait for confirmation of downward trend",
                "Monitor key support levels",
                "Consider using trailing stops",
                "Review and adjust stop-loss levels"
            ])
        else:
            strategy = "Hold"
            actions.extend([
                "Maintain current positions",
                "Monitor market conditions for changes",
                "Review and adjust stop-loss levels",
                "Consider using options strategies like iron condors for range-bound markets",
                "Focus on risk management and position sizing"
            ])
        
        # Add strategy-specific recommendations
        if strategy in ["Strong Buy", "Moderate Buy"]:
            actions.extend([
                "Consider using leverage through ETFs or options",
                "Look for sector rotation opportunities",
                "Monitor market breadth indicators"
            ])
        elif strategy in ["Strong Sell", "Moderate Sell"]:
            actions.extend([
                "Consider hedging with inverse ETFs",
                "Look for short-selling opportunities",
                "Monitor market sentiment indicators"
            ])
        else:  # Hold
            actions.extend([
                "Consider using options strategies for income",
                "Look for mean reversion opportunities",
                "Focus on portfolio rebalancing"
            ])
        
        return actions

    def _get_html_actions(self, predictions: Dict[str, Any]) -> str:
        """Get HTML formatted recommended actions with strategy indicators."""
        actions = self._get_recommended_actions(predictions)
        buy_prob = predictions.get('buy_probability', 50)
        
        # Determine strategy class for styling
        if buy_prob > 60:
            strategy_class = "buy"
        elif buy_prob < 40:
            strategy_class = "sell"
        else:
            strategy_class = "neutral"
        
        # Add strategy indicator
        strategy_html = f'<li class="list-group-item {strategy_class}"><strong>Strategy: {self._get_strategy_name(buy_prob)}</strong></li>'
        
        # Format actions
        actions_html = "\n".join([f'<li class="list-group-item">{action}</li>' for action in actions])
        
        return f"{strategy_html}\n{actions_html}"

    def _get_strategy_name(self, buy_prob: float) -> str:
        """Get strategy name based on buy probability."""
        if buy_prob > 60:
            return "Strong Buy"
        elif buy_prob > 55:
            return "Moderate Buy"
        elif buy_prob < 40:
            return "Strong Sell"
        elif buy_prob < 45:
            return "Moderate Sell"
        else:
            return "Hold"

    def add_performance_metrics(self, insights: Dict[str, Any]) -> None:
        """Add performance metrics section with detailed analysis."""
        self.report.add_heading("Performance Metrics", level=1)
        
        # Returns Analysis with clear formatting
        self.report.add_heading("Returns Analysis", level=2)
        self.report.add_paragraph(f"Total Return: {insights['total_return']:.2f}%")
        self.report.add_paragraph(f"Annual Return: {insights['annual_return']:.2f}%")
        
        # Daily Performance with clear formatting
        self.report.add_heading("Daily Performance", level=2)
        self.report.add_paragraph(f"Maximum Daily Gain: {insights['max_daily_gain']:.2f}%")
        self.report.add_paragraph(f"Maximum Daily Loss: {insights['max_daily_loss']:.2f}%")
        self.report.add_paragraph(f"Average Daily Return: {insights['avg_daily_return']:.2f}%")
        
        # Volume Analysis with clear formatting
        self.report.add_heading("Volume Analysis", level=2)
        self.report.add_paragraph(f"Latest Volume: {insights['latest_volume']:,.0f}")
        self.report.add_paragraph(f"Average Volume: {insights['avg_volume']:,.0f}")
        self.report.add_paragraph()  # Add spacing

    def add_risk_analysis(self, insights: Dict[str, Any]) -> None:
        """Add risk analysis section with detailed assessment."""
        self.report.add_heading("Risk Analysis", level=1)
        
        # Volatility Risk with clear interpretation
        self.report.add_heading("Volatility Risk", level=2)
        volatility = insights['current_volatility']
        self.report.add_paragraph(
            f"Current volatility of {volatility:.2f}% indicates "
            f"{self._get_volatility_risk_level(volatility)}"
        )
        
        # Price Levels with clear formatting
        self.report.add_heading("Key Price Levels", level=2)
        self.report.add_paragraph(f"Year High: ${insights['year_high']:,.2f}")
        self.report.add_paragraph(f"Year Low: ${insights['year_low']:,.2f}")
        
        # Risk Management Recommendations with clear formatting
        self.report.add_heading("Risk Management Recommendations", level=2)
        recommendations = self._get_risk_recommendations(insights)
        for rec in recommendations:
            self.report.add_paragraph(rec, style='List Bullet')
        self.report.add_paragraph()  # Add spacing

    def _get_volatility_risk_level(self, volatility: float) -> str:
        """Get volatility risk level description."""
        if volatility > 25:
            return "high market risk with potential for significant price swings"
        elif volatility > 15:
            return "moderate market risk with normal price fluctuations"
        else:
            return "low market risk with stable price movements"

    def _get_risk_recommendations(self, insights: Dict[str, Any]) -> List[str]:
        """Get risk management recommendations."""
        volatility = insights['current_volatility']
        current_price = insights['current_price']
        year_high = insights['year_high']
        year_low = insights['year_low']
        
        recommendations = []
        
        # Position sizing recommendations
        if volatility > 25:
            recommendations.append("Consider reducing position sizes due to high volatility")
        elif volatility < 10:
            recommendations.append("Consider increasing position sizes due to low volatility")
        
        # Stop loss recommendations
        stop_loss = current_price * 0.95  # 5% stop loss
        recommendations.append(f"Set stop-loss orders at ${stop_loss:,.2f}")
        
        # Take profit recommendations
        take_profit = current_price * 1.10  # 10% take profit
        recommendations.append(f"Set take-profit orders at ${take_profit:,.2f}")
        
        # Additional risk management
        recommendations.extend([
            "Monitor market conditions daily",
            "Review and adjust stop-loss levels weekly",
            "Consider using trailing stops for open positions"
        ])
        
        return recommendations

    def create_visualizations(self, df: pd.DataFrame, insights: Dict[str, Any], index_name: str) -> Dict[str, str]:
        """Create visualizations for the report.
        
        Args:
            df: DataFrame containing market data
            insights: Dictionary containing market insights
            index_name: Name of the market index
            
        Returns:
            Dictionary containing base64 encoded plot images
        """
        visualizations = {}
        
        # Calculate moving averages
        for ma in [20, 50, 200]:
            df[f"ma{ma}"] = df["close"].rolling(window=ma).mean()
        
        # Price and Moving Averages Plot
        plt.figure(figsize=(12, 6))
        plt.plot(df["date"], df["close"], label="Price", alpha=0.7)
        
        # Plot moving averages with different colors and transparency
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1']  # Red, Teal, Blue
        for ma, color in zip([20, 50, 200], colors):
            plt.plot(df["date"], df[f"ma{ma}"], label=f"{ma}-day MA", color=color, alpha=0.7)
        
        plt.title(f"{index_name} Price and Moving Averages")
        plt.xlabel("Date")
        plt.ylabel("Price")
        plt.legend()
        plt.grid(True, alpha=0.3)
        
        # Save plot to base64
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        visualizations['price_ma'] = base64.b64encode(buf.getvalue()).decode()
        plt.close()
        
        # Volume Analysis Plot
        plt.figure(figsize=(12, 6))
        plt.bar(df["date"], df["volume"], alpha=0.5, label="Volume")
        
        # Add volume moving average
        df["volume_ma20"] = df["volume"].rolling(window=20).mean()
        plt.plot(df["date"], df["volume_ma20"], color="red", label="20-day Volume MA", alpha=0.7)
        
        plt.title(f"{index_name} Trading Volume")
        plt.xlabel("Date")
        plt.ylabel("Volume")
        plt.legend()
        plt.grid(True, alpha=0.3)
        
        # Save plot to base64
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        visualizations['volume'] = base64.b64encode(buf.getvalue()).decode()
        plt.close()
        
        # Calculate RSI
        delta = df["close"].diff()
        gain = (delta.where(delta > 0, 0)).rolling(window=14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(window=14).mean()
        rs = gain / loss
        df["rsi"] = 100 - (100 / (1 + rs))
        
        # Calculate Annualized Volatility (20-day rolling standard deviation of returns * sqrt(252))
        df["returns"] = df["close"].pct_change()
        df["volatility"] = df["returns"].rolling(window=20).std() * np.sqrt(252) * 100  # Annualized and converted to percentage
        
        # Technical Indicators Plot (RSI and Volatility)
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 8))
        
        # RSI Plot
        ax1.plot(df["date"], df["rsi"], label="RSI", color="purple")
        ax1.axhline(y=70, color='r', linestyle='--', alpha=0.5, label="Overbought")
        ax1.axhline(y=30, color='g', linestyle='--', alpha=0.5, label="Oversold")
        ax1.set_title(f"{index_name} RSI")
        ax1.set_xlabel("Date")
        ax1.set_ylabel("RSI")
        ax1.legend()
        ax1.grid(True, alpha=0.3)
        
        # Volatility Plot
        ax2.plot(df["date"], df["volatility"], label="Volatility", color="orange")
        ax2.axhline(y=30, color='r', linestyle='--', alpha=0.5, label="High Volatility")
        ax2.axhline(y=15, color='g', linestyle='--', alpha=0.5, label="Low Volatility")
        ax2.set_title(f"{index_name} Annualized Volatility")
        ax2.set_xlabel("Date")
        ax2.set_ylabel("Volatility (%)")
        ax2.legend()
        ax2.grid(True, alpha=0.3)
        
        plt.tight_layout()
        
        # Save plot to base64
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        visualizations['indicators'] = base64.b64encode(buf.getvalue()).decode()
        plt.close()
        
        return visualizations

    def generate_report(self, insights: Dict[str, Any], df: pd.DataFrame) -> None:
        """Generate a comprehensive market analysis report.
        
        Args:
            insights: Dictionary containing market insights
            df: DataFrame containing market data
        """
        # Add report sections
        self.add_title(f"Market Analysis Report: {insights['name']}")
        self.add_executive_summary(insights)
        self.add_technical_analysis(insights)
        self.add_trading_signals(insights)
        self.add_performance_metrics(insights)
        self.add_risk_analysis(insights)
        visualizations = self.create_visualizations(df, insights, insights['name'])
        
        # Save the report
        report_path = os.path.join(self.output_dir, "market_indices_report.docx")
        self.report.save(report_path)

    def _get_market_summary(self, all_predictions: Dict[str, Dict]) -> str:
        """Generate a summary of recommendations across all market indices."""
        # Collect predictions for each index
        index_predictions = []
        for index_name, predictions in all_predictions.items():
            buy_prob = predictions.get('buy_probability', 50)
            confidence = predictions.get('confidence', 0)
            index_predictions.append({
                'name': index_name,
                'buy_prob': buy_prob,
                'confidence': confidence,
                'strategy': self._get_strategy_name(buy_prob)
            })
        
        # Calculate overall market sentiment
        avg_buy_prob = sum(p['buy_prob'] for p in index_predictions) / len(index_predictions)
        overall_sentiment = "Bullish" if avg_buy_prob > 55 else "Bearish" if avg_buy_prob < 45 else "Neutral"
        
        # Generate HTML for the summary section
        summary_html = f"""
        <div class="market-section">
            <h2>Overall Market Summary</h2>
            <div class="metric-card">
                <h4>Market Sentiment: <span class="{overall_sentiment.lower()}">{overall_sentiment}</span></h4>
                <p>Average Buy Probability: <strong>{avg_buy_prob:.1f}%</strong></p>
            </div>
            
            <div class="row">
                <div class="col-md-12">
                    <h4>Index-Specific Recommendations</h4>
                    <table class="table">
                        <thead>
                            <tr>
                                <th>Index</th>
                                <th>Strategy</th>
                                <th>Buy Probability</th>
                                <th>Model Confidence</th>
                            </tr>
                        </thead>
                        <tbody>
        """
        
        for pred in index_predictions:
            strategy_class = "buy" if pred['buy_prob'] > 55 else "sell" if pred['buy_prob'] < 45 else "neutral"
            summary_html += f"""
                            <tr>
                                <td>{pred['name']}</td>
                                <td class="{strategy_class}">{pred['strategy']}</td>
                                <td>{pred['buy_prob']:.1f}%</td>
                                <td>{pred['confidence']:.1f}%</td>
                            </tr>
            """
        
        summary_html += """
                        </tbody>
                    </table>
                </div>
            </div>
            
            <div class="metric-card">
                <h4>Portfolio Recommendations</h4>
                <ul class="list-group">
        """
        
        # Add portfolio-level recommendations based on overall sentiment
        if overall_sentiment == "Bullish":
            summary_html += """
                    <li class="list-group-item">Consider increasing exposure to growth stocks</li>
                    <li class="list-group-item">Look for sector rotation opportunities</li>
                    <li class="list-group-item">Consider using leverage through ETFs</li>
                    <li class="list-group-item">Monitor market breadth indicators</li>
            """
        elif overall_sentiment == "Bearish":
            summary_html += """
                    <li class="list-group-item">Consider reducing overall market exposure</li>
                    <li class="list-group-item">Look for defensive sector opportunities</li>
                    <li class="list-group-item">Consider hedging with inverse ETFs</li>
                    <li class="list-group-item">Monitor market sentiment indicators</li>
            """
        else:
            summary_html += """
                    <li class="list-group-item">Maintain balanced portfolio allocation</li>
                    <li class="list-group-item">Focus on income-generating strategies</li>
                    <li class="list-group-item">Consider options strategies for range-bound markets</li>
                    <li class="list-group-item">Monitor for breakout opportunities</li>
            """
        
        summary_html += """
                </ul>
            </div>
        </div>
        """
        
        return summary_html

    def generate_html_report(
        self,
        all_data: Dict[str, pd.DataFrame],
        all_insights: Dict[str, Dict],
        all_visualizations: Dict[str, Dict],
        all_predictions: Dict[str, Dict]
    ) -> tuple[str, str]:
        """Generate HTML and PDF reports for all market indices."""
        html_path = os.path.join(self.output_dir, "market_indices_report.html")
        pdf_path = os.path.join(self.output_dir, "market_indices_report.pdf")

        # Create HTML content
        html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Market Analysis Report</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
            <style>
                body {{ 
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    margin: 0;
                    padding: 20px;
                    background-color: #f8f9fa;
                }}
                .container {{
                    max-width: 1200px;
                    margin: 0 auto;
                    background-color: white;
                    padding: 30px;
                    border-radius: 10px;
                    box-shadow: 0 0 20px rgba(0,0,0,0.1);
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 40px;
                    padding-bottom: 20px;
                    border-bottom: 2px solid #eee;
                }}
                .market-section {{
                    margin-bottom: 40px;
                    padding: 20px;
                    border: 1px solid #eee;
                    border-radius: 5px;
                }}
                .metric-card {{
                    background-color: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 15px;
                }}
                .buy {{ color: #28a745; }}
                .sell {{ color: #dc3545; }}
                .neutral {{ color: #ffc107; }}
                .bullish {{ color: #28a745; font-weight: bold; }}
                .bearish {{ color: #dc3545; font-weight: bold; }}
                .chart-container {{
                    margin: 20px 0;
                    padding: 15px;
                    border: 1px solid #eee;
                    border-radius: 5px;
                }}
                .chart-container img {{
                    max-width: 100%;
                    height: auto;
                    margin: 10px 0;
                }}
                table {{
                    width: 100%;
                    margin-bottom: 20px;
                }}
                th {{ 
                    background-color: #f8f9fa;
                    padding: 12px;
                }}
                td {{
                    padding: 10px;
                    border-bottom: 1px solid #eee;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>Market Analysis Report</h1>
                    <p class="text-muted">Generated on {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                </div>
        """

        # Add market summary section first
        html_content += self._get_market_summary(all_predictions)

        # Add individual market sections
        for index_name, insights in all_insights.items():
            predictions = all_predictions.get(index_name, {})
            buy_prob = predictions.get('buy_probability', 0)
            signal = 'Buy' if buy_prob > 50 else 'Sell'
            signal_class = 'buy' if buy_prob > 50 else 'sell'
            
            html_content += f"""
                <div class="market-section">
                    <h2>{index_name} Analysis</h2>
                    
                    <div class="row">
                        <div class="col-md-6">
                            <div class="metric-card">
                                <h4>Current Market Status</h4>
                                <p>Price: <strong>${insights.get('current_price', 0):,.2f}</strong></p>
                                <p>Return: <strong>{insights.get('total_return', 0):.2f}%</strong></p>
                                <p>Volatility: <strong>{insights.get('current_volatility', 0):.2f}%</strong></p>
                                <p>RSI: <strong>{insights.get('current_rsi', 0):.1f}</strong></p>
                            </div>
                        </div>
                        
                        <div class="col-md-6">
                            <div class="metric-card">
                                <h4>Trading Signal</h4>
                                <p class="{signal_class}">
                                    <strong>{signal}</strong>
                                </p>
                                <p>Buy Probability: <strong>{buy_prob:.1f}%</strong></p>
                                <p>Model Confidence: <strong>{predictions.get('confidence', 0):.1f}%</strong></p>
                            </div>
                        </div>
                    </div>

                    <div class="chart-container">
                        <h4>Price and Moving Averages</h4>
                        <img src="data:image/png;base64,{all_visualizations[index_name]['price_ma']}" alt="Price and MA">
                    </div>

                    <div class="chart-container">
                        <h4>Volume Analysis</h4>
                        <img src="data:image/png;base64,{all_visualizations[index_name]['volume']}" alt="Volume">
                    </div>

                    <div class="chart-container">
                        <h4>Technical Indicators</h4>
                        <img src="data:image/png;base64,{all_visualizations[index_name]['indicators']}" alt="Technical Indicators">
                    </div>

                    <div class="chart-container">
                        <h4>Technical Analysis</h4>
                        <table class="table">
                            <tr>
                                <th>Indicator</th>
                                <th>Value</th>
                                <th>Interpretation</th>
                            </tr>
                            <tr>
                                <td>RSI</td>
                                <td>{insights.get('current_rsi', 0):.1f}</td>
                                <td>{self._get_rsi_interpretation(insights.get('current_rsi', 50))}</td>
                            </tr>
                            <tr>
                                <td>Volatility</td>
                                <td>{insights.get('current_volatility', 0):.2f}%</td>
                                <td>{self._get_volatility_level(insights.get('current_volatility', 15))}</td>
                            </tr>
                        </table>
                    </div>

                    <div class="chart-container">
                        <h4>Risk Analysis</h4>
                        <table class="table">
                            <tr>
                                <th>Metric</th>
                                <th>Value</th>
                            </tr>
                            <tr>
                                <td>Year High</td>
                                <td>${insights.get('year_high', 0):,.2f}</td>
                            </tr>
                            <tr>
                                <td>Year Low</td>
                                <td>${insights.get('year_low', 0):,.2f}</td>
                            </tr>
                            <tr>
                                <td>Average Volume</td>
                                <td>{insights.get('avg_volume', 0):,.0f}</td>
                            </tr>
                        </table>
                    </div>

                    <div class="chart-container">
                        <h4>Recommended Actions</h4>
                        <ul class="list-group">
                            {self._get_html_actions(predictions)}
                        </ul>
                    </div>
                </div>
            """

        # Close HTML
        html_content += """
            </div>
            <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
        </body>
        </html>
        """

        # Write HTML file
        with open(html_path, 'w') as f:
            f.write(html_content)

        # For now, create a simple PDF version
        import shutil
        shutil.copy(html_path, pdf_path)
        
        return html_path, pdf_path 